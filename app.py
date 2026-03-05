import datetime
import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from io import BytesIO

# --- ИНТЕРФЕЙС ---
st.set_page_config(page_title="Генератор отчетов", page_icon="🖨")
st.title("Генератор отчетов по успеваемости")

# --- БОКОВАЯ ПАНЕЛЬ НАСТРОЕК ---
st.sidebar.header("⚙️ Настройки отчета")

# Настройка фильтра по дате
st.sidebar.subheader("Фильтр данных")
start_date = st.sidebar.date_input(
    "Показывать оценки начиная с:",
    value=datetime.date(2026, 1, 1),
    help="Оценки за даты ранее выбранной не попадут в обратную связь"
)

# Настройка внешнего вида
st.sidebar.subheader("Внешний вид")
compact_mode = st.sidebar.toggle("Компактный режим (Beta)", value=False,
                                 help="Размещает таблицы в ряд для экономии бумаги")

orientation = st.sidebar.selectbox("Ориентация страницы", ["Книжная", "Альбомная"])

show_dates = st.sidebar.toggle("Включить строку дат", value=True)

separator_type = st.sidebar.selectbox(
    "Разделитель между учениками",
    ["Разрыв страницы", "Пустые строки (параграфы)"]
)

num_paragraphs = 1
if separator_type == "Пустые строки (параграфы)":
    num_paragraphs = st.sidebar.number_input("Количество пустых строк", min_value=1, max_value=10, value=2)

# КОНСТАНТЫ ТАБЛИЦЫ
MAX_WIDTH_CM = 2
BASE_HEIGHT_CM = 1.5
HEIGHT_COEFF = 0.07
MAX_TOPIC_HEIGHT_CM = 5.5

# Динамическая ширина страницы в зависимости от ориентации
if orientation == "Альбомная":
    PAGE_WIDTH_CM = 27.5  # ~29.7см ширина листа минус поля
else:
    PAGE_WIDTH_CM = 19.0  # ~21см ширина листа минус поля


def cm_to_dxa(cm):
    inches = cm / 2.54
    points = inches * 72
    return int(round(points * 20))


def format_grade(val):
    if pd.isna(val): return ""
    if isinstance(val, (datetime.datetime, datetime.date, pd.Timestamp)):
        return val.strftime("%d/%m").lstrip("0").replace("/0", "/")
    return str(val).strip()


max_col_width_dxa = cm_to_dxa(MAX_WIDTH_CM)

uploaded_file = st.file_uploader("Выберите Excel файл (журнал класса)", type=["xlsx"])

if uploaded_file:
    if st.button("🚀 Создать ОС"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        output_doc = BytesIO()
        results = []

        try:
            xls = pd.ExcelFile(uploaded_file)
            sheet_names = xls.sheet_names
            index_df = pd.read_excel(uploaded_file, sheet_name=sheet_names[0], header=None)
            subject_sheets = index_df.iloc[:, 0].dropna().tolist()

            # ------------------ ЭТАП 1: СБОР ДАННЫХ ------------------
            for i, sheet in enumerate(subject_sheets):
                progress_bar.progress(i / len(subject_sheets))
                status_text.text(f"🔍 Обработка: {sheet}")
                if sheet not in sheet_names: continue
                df = pd.read_excel(uploaded_file, sheet_name=sheet, header=None)

                try:
                    topics, dates_raw, students = df.iloc[0, 3:], df.iloc[1, 3:], df.iloc[5:, :]
                    for _, row in students.iterrows():
                        student = row[1]
                        if not isinstance(student, str) or not student.strip(): continue
                        for t, d, g in zip(topics, dates_raw, row[3:]):
                            try:
                                c_date = pd.to_datetime(d).date()
                                if c_date < start_date: continue
                                formatted_g = format_grade(g)
                                if not formatted_g or formatted_g.lower() == "nan": continue
                                results.append({
                                    "ФИО": student.strip(),
                                    "Предмет": str(sheet).strip(),
                                    "Тема": str(t).strip(),
                                    "Дата": c_date,
                                    "Оценка": formatted_g
                                })
                            except:
                                continue
                except:
                    continue

            if not results:
                st.error("Оценок не найдено.")
            else:
                with st.spinner("Формирование документа..."):
                    doc = Document()

                    # Настройка ориентации и полей
                    section = doc.sections[0]
                    if orientation == "Альбомная":
                        section.orientation = WD_ORIENT.LANDSCAPE
                        section.page_width, section.page_height = Cm(29.7), Cm(21.0)

                    margin = 0.8 if compact_mode else 1.0
                    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(margin)

                    full = pd.DataFrame(results).sort_values(["ФИО", "Предмет", "Дата"])
                    student_groups = list(full.groupby("ФИО"))

                    for idx, (student, df_student) in enumerate(student_groups):
                        status_text.text(f"📄 Оформление: {student}")
                        doc.add_heading(student, level=2)

                        subjs = []
                        for s_name, df_s in df_student.groupby("Предмет"):
                            t_list = df_s["Тема"].tolist()
                            formatted_dates = [d.strftime("%d.%m") for d in df_s["Дата"]]
                            subjs.append({
                                "name": s_name,
                                "topics": t_list,
                                "dates": formatted_dates,
                                "grades": df_s["Оценка"].tolist(),
                                "w": len(t_list) * MAX_WIDTH_CM
                            })

                        if compact_mode:
                            # --- КОМПАКТНЫЙ РЕЖИМ (ПРОПОРЦИОНАЛЬНЫЙ) ---
                            rows, curr_row, curr_w = [], [], 0
                            for s in subjs:
                                if curr_w + s["w"] > PAGE_WIDTH_CM and curr_row:
                                    rows.append(curr_row)
                                    curr_row, curr_w = [], 0
                                curr_row.append(s);
                                curr_w += s["w"] + 0.5
                            if curr_row: rows.append(curr_row)

                            for r_subjs in rows:
                                total_cols = sum(len(s["topics"]) for s in r_subjs)
                                container = doc.add_table(rows=1, cols=len(r_subjs))
                                container.autofit = False

                                for c_idx, s in enumerate(r_subjs):
                                    cell = container.rows[0].cells[c_idx]
                                    # Пропорциональная ширина ячейки контейнера
                                    c_width = int((len(s["topics"]) / total_cols) * cm_to_dxa(PAGE_WIDTH_CM))
                                    cell._tc.get_or_add_tcPr().get_or_add_tcW().set(qn('w:w'), str(c_width))

                                    cell.paragraphs[0].add_run(s["name"]).bold = True
                                    n_rows = 3 if show_dates else 2
                                    inner = cell.add_table(rows=n_rows, cols=len(s["topics"]))
                                    inner.style = 'Table Grid'

                                    # Фиксация колонок внутри
                                    inner._tbl.tblPr.append(
                                        parse_xml(r'<w:tblLayout {} w:type="fixed"/>'.format(nsdecls('w'))))

                                    for ci, (t, d, g) in enumerate(zip(s["topics"], s["dates"], s["grades"])):
                                        inner.rows[0].cells[ci].text = str(t)
                                        tcPr = inner.rows[0].cells[ci]._tc.get_or_add_tcPr()
                                        tcPr.append(
                                            parse_xml(r'<w:textDirection {} w:val="btLr"/>'.format(nsdecls('w'))))
                                        if show_dates: inner.rows[1].cells[ci].text = str(d)
                                        inner.rows[-1].cells[ci].text = str(g)
                                        for ri in range(n_rows):
                                            itcW = inner.rows[ri].cells[ci]._tc.get_or_add_tcPr().get_or_add_tcW()
                                            itcW.set(qn('w:w'), str(max_col_width_dxa))

                                    calculated_h = max(len(str(t)) for t in s["topics"]) * HEIGHT_COEFF
                                    final_h = max(BASE_HEIGHT_CM, min(calculated_h, MAX_TOPIC_HEIGHT_CM))
                                    inner.rows[0].height = Cm(final_h)
                                    inner.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                                doc.add_paragraph()
                        else:
                            # --- КЛАССИЧЕСКИЙ РЕЖИМ ---
                            for s in subjs:
                                doc.add_heading(s["name"], level=3)
                                n_rows = 3 if show_dates else 2
                                table = doc.add_table(rows=n_rows, cols=len(s["topics"]))
                                table.style = 'Table Grid'
                                for ci, (t, d, g) in enumerate(zip(s["topics"], s["dates"], s["grades"])):
                                    table.rows[0].cells[ci].text = t
                                    tcPr = table.rows[0].cells[ci]._tc.get_or_add_tcPr()
                                    tcPr.append(parse_xml(r'<w:textDirection {} w:val="btLr"/>'.format(nsdecls('w'))))
                                    if show_dates: table.rows[1].cells[ci].text = d
                                    table.rows[-1].cells[ci].text = g
                                    for ri in range(n_rows):
                                        table.rows[ri].cells[ci]._tc.get_or_add_tcPr().get_or_add_tcW().set(qn('w:w'),
                                                                                                            str(max_col_width_dxa))
                                table.rows[0].height = Cm(
                                    max(max(len(str(t)) for t in s["topics"]) * HEIGHT_COEFF, BASE_HEIGHT_CM))

                        if idx < len(student_groups) - 1:
                            if separator_type == "Разрыв страницы":
                                doc.add_page_break()
                            else:
                                [doc.add_paragraph() for _ in range(num_paragraphs)]

                doc.save(output_doc)
                progress_bar.progress(1.0)
                status_text.empty()
                st.success("✅ Готово!")

                filename = (f"Обратная_Связь_{datetime.datetime.now().strftime('%d.%m')} "
                            f"({uploaded_file.name.split('.')[0]}).docx")
                st.download_button(
                    label="📥 Скачать обратную связь",
                    data=output_doc.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Критическая ошибка: {e}")